"""
Handler for .docx files containing HL7 messages in paragraph text or tables.
"""
import io
import re
from typing import List

from app.file_handlers.base import BaseFileHandler


class DocxHandler(BaseFileHandler):

    def can_handle(self, filename: str, content_type: str) -> bool:
        ext = "." + filename.rsplit(".", 1)[-1].lower() if "." in filename else ""
        return ext == ".docx"

    def extract_messages(self, data: bytes, filename: str) -> List[str]:
        try:
            from docx import Document
        except ImportError:
            raise ValueError("python-docx is required to read .docx files.")

        try:
            doc = Document(io.BytesIO(data))
        except Exception as exc:
            raise ValueError(f"Could not open .docx file: {exc}") from exc

        # Collect all text blocks: paragraphs and table cells
        text_blocks = []
        for para in doc.paragraphs:
            text = para.text.strip()
            if text:
                text_blocks.append(text)
        for table in doc.tables:
            for row in table.rows:
                for cell in row.cells:
                    text = cell.text.strip()
                    if text:
                        text_blocks.append(text)

        # Join all text and split on MSH boundaries
        full_text = "\n".join(text_blocks)
        full_text = full_text.replace("\r\n", "\r").replace("\n", "\r")
        parts = re.split(r"(?=MSH\|)", full_text, flags=re.IGNORECASE)
        messages = [p.strip() for p in parts if p.strip().upper().startswith("MSH|")]

        if not messages:
            raise ValueError(
                "No HL7 messages found in .docx file. "
                "The document should contain HL7 message text starting with MSH|."
            )
        return messages
